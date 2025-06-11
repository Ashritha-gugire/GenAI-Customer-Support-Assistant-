#!/usr/bin/env python
# coding: utf-8

import streamlit as st
import os
from datetime import datetime
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go

# LangChain imports (updated)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings

from langchain_openai import OpenAIEmbeddings, ChatOpenAI
#from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory

# Document processing
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument
import io
import json
import uuid
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class RAGCustomerSupport:
    def __init__(self):
        self.setup_components()
        self.conversation_history = []
        self.analytics_data = {
            'queries': [],
            'response_times': [],
            'satisfaction_scores': [],
            'topics': []
        }
    
    def setup_components(self):
        """Initialize RAG components"""
        try:
            # Check for OpenAI API key
            if 'openai_api_key' not in st.session_state:
                st.session_state.openai_api_key = ""
            
            # Initialize embeddings (using Hugging Face as fallback)
            if st.session_state.openai_api_key and st.session_state.openai_api_key.startswith('sk-'):
                os.environ["OPENAI_API_KEY"] = st.session_state.openai_api_key
                try:
                    self.embeddings = OpenAIEmbeddings()
                    self.llm = ChatOpenAI(temperature=0.7, model="gpt-3.5-turbo")
                    logger.info("Using OpenAI embeddings and LLM")
                except Exception as e:
                    logger.error(f"Error initializing OpenAI: {e}")
                    self.setup_fallback()
            else:
                self.setup_fallback()
            
            # Text splitter for chunking documents
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            )
            
            # Initialize vector store
            self.vector_store = None
            self.qa_chain = None
            
            # Memory for conversation context
            self.memory = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True
            )
            
            # Default knowledge base
            self.load_default_knowledge()
            
        except Exception as e:
            logger.error(f"Error initializing components: {e}")
            st.error(f"Error initializing components: {e}")
    
    def setup_fallback(self):
        """Setup fallback embeddings when OpenAI is not available"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2"
            )
            self.llm = None
            logger.info("Using HuggingFace embeddings with fallback response generation")
        except Exception as e:
            logger.error(f"Error setting up fallback: {e}")
            st.error(f"Error setting up fallback embeddings: {e}")
    
    def load_default_knowledge(self):
        """Load default knowledge base documents"""
        default_docs = [
            {
                "title": "Return Policy",
                "content": """
                Our Return Policy:
                
                We offer a 30-day return policy for all products purchased from our store.
                
                Eligibility:
                - Items must be returned within 30 days of purchase
                - Items must be in original condition with all packaging
                - Items must include original receipt or order number
                
                Process:
                1. Contact customer service with your order number
                2. Receive return authorization and shipping label
                3. Package item securely and ship using provided label
                4. Refund will be processed within 5-7 business days
                
                Exceptions:
                - Custom or personalized items cannot be returned
                - Digital products are non-refundable
                - Sale items may have different return conditions
                
                Contact Information:
                - Email: returns@company.com
                - Phone: 1-800-RETURNS
                - Live chat available Monday-Friday 9 AM - 6 PM EST
                """,
                "category": "Returns"
            },
            {
                "title": "Shipping Information",
                "content": """
                Shipping Options and Information:
                
                Standard Shipping:
                - Delivery time: 3-5 business days
                - Cost: $5.99 (Free on orders over $50)
                
                Express Shipping:
                - Delivery time: 1-2 business days
                - Cost: $12.99
                
                Overnight Shipping:
                - Delivery time: Next business day
                - Cost: $24.99
                
                International Shipping:
                - Available to most countries
                - Delivery time: 7-14 business days
                - Costs vary by destination
                
                Order Processing:
                - Orders placed before 2 PM EST ship same day
                - Weekend orders ship on Monday
                - Tracking information provided via email
                
                Shipping Restrictions:
                - Some items cannot be shipped to PO boxes
                - Hazardous materials have special shipping requirements
                - International orders may be subject to customs fees
                """,
                "category": "Shipping"
            },
            {
                "title": "Account Management",
                "content": """
                Managing Your Account:
                
                Creating an Account:
                - Visit our website and click 'Sign Up'
                - Provide email address and create password
                - Verify email address to activate account
                
                Account Features:
                - View order history and tracking
                - Save multiple shipping addresses
                - Store payment methods securely
                - Manage email preferences
                - Track loyalty points and rewards
                
                Password Reset:
                - Click 'Forgot Password' on login page
                - Enter your email address
                - Check email for reset link
                - Create new password
                
                Updating Information:
                - Log into your account
                - Navigate to 'Account Settings'
                - Update personal information, addresses, or payment methods
                
                Account Security:
                - Use strong, unique passwords
                - Enable two-factor authentication
                - Monitor account activity regularly
                - Contact us immediately if you notice suspicious activity
                """,
                "category": "Account"
            },
            {
                "title": "Technical Support",
                "content": """
                Technical Support Guide:
                
                Common Issues and Solutions:
                
                Device Won't Turn On:
                1. Check power connection and cables
                2. Try different power outlet
                3. Hold power button for 10 seconds to reset
                4. Contact support if issue persists
                
                Connectivity Problems:
                1. Check WiFi connection
                2. Restart your router
                3. Move device closer to router
                4. Check for software updates
                
                App Issues:
                1. Close and restart the app
                2. Check for app updates in store
                3. Clear app cache and data
                4. Reinstall app if necessary
                
                Performance Issues:
                1. Close unnecessary applications
                2. Restart your device
                3. Check available storage space
                4. Update device software
                
                Getting Help:
                - Check our online knowledge base
                - Submit support ticket with device model and issue description
                - Live chat available Monday-Friday 9 AM - 6 PM EST
                - Phone support: 1-800-SUPPORT
                - Video tutorials available on our YouTube channel
                """,
                "category": "Technical"
            },
            {
                "title": "Payment and Billing",
                "content": """
                Payment Information:
                
                Accepted Payment Methods:
                - Visa, MasterCard, American Express, Discover
                - PayPal and PayPal Credit
                - Apple Pay and Google Pay
                - Buy now, pay later options (Klarna, Afterpay)
                - Gift cards and store credit
                
                Security:
                - All payments processed through secure SSL encryption
                - Credit card information is not stored on our servers
                - PCI DSS compliant payment processing
                
                Billing:
                - Charges appear as 'TechStore Inc' on statements
                - Billing occurs when items ship
                - Automatic payment retry for failed transactions
                
                Billing Issues:
                - Incorrect charges will be investigated within 24 hours
                - Refunds processed within 5-7 business days
                - Payment failures may delay order processing
                
                Subscription Services:
                - Monthly or annual billing options
                - Cancel anytime through account settings
                - Prorated refunds for early cancellation
                - Automatic renewal notifications sent 7 days before billing
                """,
                "category": "Payment"
            }
        ]
        
        # Convert to Document objects
        documents = []
        for doc in default_docs:
            documents.append(Document(
                page_content=doc["content"],
                metadata={
                    "title": doc["title"],
                    "category": doc["category"],
                    "source": "default_kb"
                }
            ))
        
        # Create vector store
        self.create_vector_store(documents)
    
    def create_vector_store(self, documents):
        """Create or update vector store with documents"""
        try:
            if not self.embeddings:
                logger.error("Embeddings not initialized")
                return False
                
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            logger.info(f"Split documents into {len(texts)} chunks")
            
            # Create vector store
            if self.vector_store is None:
                self.vector_store = FAISS.from_documents(texts, self.embeddings)
                logger.info("Created new FAISS vector store")
            else:
                # Add new documents to existing store
                self.vector_store.add_documents(texts)
                logger.info("Added documents to existing vector store")
            
            # Create QA chain if LLM is available
            if self.llm:
                self.create_qa_chain()
            
            return True
            
        except Exception as e:
            logger.error(f"Error creating vector store: {e}")
            st.error(f"Error creating vector store: {e}")
            return False
    
    def create_qa_chain(self):
        """Create the question-answering chain"""
        try:
            # Custom prompt template
            prompt_template = """
            You are a helpful customer support assistant. Use the following context to answer the customer's question.
            If you don't know the answer based on the context, say so and suggest contacting human support.
            
            Context: {context}
            
            Question: {question}
            
            Please provide a helpful, accurate, and friendly response. Include relevant details from the context.
            If applicable, mention next steps the customer should take.
            
            Answer:
            """
            
            PROMPT = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"]
            )
            
            # Create retrieval QA chain
            self.qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=self.vector_store.as_retriever(search_kwargs={"k": 3}),
                chain_type_kwargs={"prompt": PROMPT},
                return_source_documents=True
            )
            logger.info("Created QA chain successfully")
            
        except Exception as e:
            logger.error(f"Error creating QA chain: {e}")
            st.error(f"Error creating QA chain: {e}")
    
    def process_uploaded_file(self, uploaded_file):
        """Process uploaded document and add to knowledge base"""
        try:
            # Read file content based on type
            if uploaded_file.type == "text/plain":
                content = str(uploaded_file.read(), "utf-8")
            elif uploaded_file.type == "application/pdf":
                # Process PDF
                pdf_bytes = uploaded_file.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                content = ""
                for page in doc:
                    content += page.get_text()
                doc.close()
            elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                # Process DOCX
                docx_bytes = uploaded_file.read()
                doc = DocxDocument(io.BytesIO(docx_bytes))
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            else:
                st.error(f"Unsupported file type: {uploaded_file.type}")
                return False
            
            if not content.strip():
                st.error("File appears to be empty or could not be processed.")
                return False
            
            # Create document object
            document = Document(
                page_content=content,
                metadata={
                    "title": uploaded_file.name,
                    "category": "Uploaded",
                    "source": "user_upload",
                    "upload_time": datetime.now().isoformat()
                }
            )
            
            # Add to vector store
            success = self.create_vector_store([document])
            
            if success:
                st.session_state.uploaded_docs.append({
                    "name": uploaded_file.name,
                    "size": uploaded_file.size,
                    "type": uploaded_file.type,
                    "upload_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
                logger.info(f"Successfully processed file: {uploaded_file.name}")
            
            return success
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {e}")
            st.error(f"Error processing file {uploaded_file.name}: {e}")
            return False
    
    def get_response(self, query):
        """Get response to user query using RAG"""
        start_time = datetime.now()
        
        try:
            if self.qa_chain and self.llm:
                # Use OpenAI-powered RAG
                result = self.qa_chain({"query": query})
                response = result["result"]
                source_docs = result.get("source_documents", [])
                
                # Extract source information
                sources = []
                for doc in source_docs:
                    sources.append({
                        "title": doc.metadata.get("title", "Unknown"),
                        "category": doc.metadata.get("category", "Unknown"),
                        "content_snippet": doc.page_content[:200] + "..."
                    })
                
            else:
                # Fallback: Simple retrieval without LLM
                if self.vector_store:
                    docs = self.vector_store.similarity_search(query, k=3)
                    
                    if docs:
                        # Enhanced rule-based response generation
                        response = self.generate_enhanced_response(query, docs)
                        sources = [{
                            "title": doc.metadata.get("title", "Unknown"),
                            "category": doc.metadata.get("category", "Unknown"),
                            "content_snippet": doc.page_content[:200] + "..."
                        } for doc in docs]
                    else:
                        response = "I don't have information about that topic in my knowledge base. Please contact our support team for assistance."
                        sources = []
                else:
                    response = "Knowledge base not available. Please contact our support team."
                    sources = []
            
            # Calculate response time
            response_time = (datetime.now() - start_time).total_seconds()
            
            # Log analytics
            self.analytics_data['queries'].append(query)
            self.analytics_data['response_times'].append(response_time)
            self.analytics_data['topics'].append(self.categorize_query(query))
            
            logger.info(f"Generated response for query: {query[:50]}... (Time: {response_time:.2f}s)")
            
            return {
                "response": response,
                "sources": sources,
                "response_time": response_time
            }
            
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            return {
                "response": "I encountered an error processing your request. Please try again or contact support.",
                "sources": [],
                "response_time": 0
            }
    
    def generate_enhanced_response(self, query, docs):
        """Generate enhanced response without LLM using better logic"""
        query_lower = query.lower()
        
        # Extract the most relevant content from documents
        relevant_content = ""
        for doc in docs[:2]:  # Use top 2 most relevant docs
            relevant_content += doc.page_content[:400] + "\n\n"
        
        # Enhanced keyword matching for response generation
        if any(word in query_lower for word in ['return', 'refund', 'exchange', 'send back']):
            return f"Here's information about returns and refunds:\n\n{relevant_content}\n\nIf you need to start a return, please contact our customer service team with your order number. They'll provide you with a return authorization and shipping label."
        
        elif any(word in query_lower for word in ['shipping', 'delivery', 'ship', 'track', 'when will']):
            return f"Here's our shipping information:\n\n{relevant_content}\n\nFor specific tracking information, please check your email for tracking details or contact our support team with your order number."
        
        elif any(word in query_lower for word in ['account', 'login', 'password', 'sign in', 'forgot']):
            return f"Here's help with account management:\n\n{relevant_content}\n\nIf you're still having trouble with your account, please try the password reset option or contact our support team for personalized assistance."
        
        elif any(word in query_lower for word in ['payment', 'billing', 'charge', 'credit card', 'pay']):
            return f"Here's information about payment and billing:\n\n{relevant_content}\n\nFor specific billing questions or disputes, please contact our billing department directly."
        
        elif any(word in query_lower for word in ['technical', 'problem', 'issue', 'broken', 'not working', 'error']):
            return f"Here's technical support information:\n\n{relevant_content}\n\nIf these steps don't resolve your issue, please contact our technical support team with details about your device model and the specific problem you're experiencing."
        
        elif any(word in query_lower for word in ['how', 'what', 'when', 'where', 'why']):
            return f"Based on our documentation:\n\n{relevant_content}\n\nIs there anything specific about this information you'd like me to clarify? Feel free to ask a more specific question or contact our support team for detailed assistance."
        
        else:
            return f"Here's relevant information from our knowledge base:\n\n{relevant_content}\n\nIf this doesn't fully answer your question, please feel free to ask for more specific information or contact our support team for personalized assistance."
    
    def categorize_query(self, query):
        """Categorize user query for analytics"""
        query_lower = query.lower()
        
        if any(word in query_lower for word in ['return', 'refund', 'exchange']):
            return 'Returns'
        elif any(word in query_lower for word in ['shipping', 'delivery']):
            return 'Shipping'
        elif any(word in query_lower for word in ['account', 'login']):
            return 'Account'
        elif any(word in query_lower for word in ['payment', 'billing']):
            return 'Payment'
        elif any(word in query_lower for word in ['technical', 'problem']):
            return 'Technical'
        else:
            return 'General'


def main():
    st.set_page_config(
        page_title="AI Customer Support",
        page_icon="🤖",
        layout="wide"
    )
    
    st.title("GenAI Customer Support Assistant")
    st.markdown("RAG-powered AI assistant with document upload and search capabilities")
    
    # Initialize session state
    if 'rag_system' not in st.session_state:
        with st.spinner("Initializing AI system..."):
            st.session_state.rag_system = RAGCustomerSupport()
    
    if 'messages' not in st.session_state:
        st.session_state.messages = [{
            "role": "assistant",
            "content": "Hello! I'm your AI Customer Support Assistant. I can help you with questions about returns, shipping, accounts, technical issues, and more. You can also upload documents to expand my knowledge base!",
            "timestamp": datetime.now().strftime("%H:%M:%S")
        }]
    
    if 'uploaded_docs' not in st.session_state:
        st.session_state.uploaded_docs = []
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key input
        api_key = st.text_input(
            "OpenAI API Key (Optional)",
            type="password",
            value=st.session_state.openai_api_key,
            help="Enter your OpenAI API key for enhanced responses. Leave blank to use open-source models."
        )
        
        if api_key != st.session_state.openai_api_key:
            st.session_state.openai_api_key = api_key
            # Reinitialize system with new API key
            with st.spinner("Updating AI system..."):
                st.session_state.rag_system = RAGCustomerSupport()
            st.rerun()
        
        # Show current mode
        if st.session_state.openai_api_key and st.session_state.openai_api_key.startswith('sk-'):
            st.success(" OpenAI Mode Active")
        else:
            st.info("Open Source Mode Active")
        
        st.divider()
        
        # File upload
        st.subheader(" Upload Documents")
        uploaded_file = st.file_uploader(
            "Upload knowledge base documents",
            type=['txt', 'pdf', 'docx'],
            help="Upload documents to expand the AI's knowledge base"
        )
        
        if uploaded_file and st.button("Add to Knowledge Base"):
            with st.spinner("Processing document..."):
                success = st.session_state.rag_system.process_uploaded_file(uploaded_file)
                if success:
                    st.success(f" Successfully added {uploaded_file.name} to knowledge base!")
                    st.rerun()
        
        st.divider()
        
        # Uploaded documents
        if st.session_state.uploaded_docs:
            st.subheader("Uploaded Documents")
            for doc in st.session_state.uploaded_docs:
                with st.expander(f"📎 {doc['name']}"):
                    st.write(f"**Size:** {doc['size']} bytes")
                    st.write(f"**Type:** {doc['type']}")
                    st.write(f"**Uploaded:** {doc['upload_time']}")
        
        st.divider()
        
        # Quick actions
        st.subheader("Quick Actions")
        if st.button("Clear Chat History"):
            st.session_state.messages = st.session_state.messages[:1]  # Keep welcome message
            st.rerun()
        
        if st.button("Download Chat Log"):
            chat_log = "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.messages])
            st.download_button(
                label="Download",
                data=chat_log,
                file_name=f"chat_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )
    
    # Main chat interface
    tab1, tab2, tab3 = st.tabs(["💬 Chat", "Analytics", "Knowledge Base"])
    
    with tab1:
        # Display chat messages
        chat_container = st.container()
        with chat_container:
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.write(message["content"])
                    if "sources" in message and message["sources"]:
                        with st.expander("Sources"):
                            for i, source in enumerate(message["sources"]):
                                st.write(f"**{i+1}. {source['title']}** ({source['category']})")
                                st.write(source['content_snippet'])
                    if "timestamp" in message:
                        st.caption(f"{message['timestamp']}")
        
        # Chat input
        if prompt := st.chat_input("Ask me anything about our products or services..."):
            # Add user message
            timestamp = datetime.now().strftime("%H:%M:%S")
            st.session_state.messages.append({
                "role": "user",
                "content": prompt,
                "timestamp": timestamp
            })
            
            # Display user message
            with st.chat_message("user"):
                st.write(prompt)
                st.caption(f"{timestamp}")
            
            # Get AI response
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    result = st.session_state.rag_system.get_response(prompt)
                
                st.write(result["response"])
                
                if result["sources"]:
                    with st.expander("Sources"):
                        for i, source in enumerate(result["sources"]):
                            st.write(f"**{i+1}. {source['title']}** ({source['category']})")
                            st.write(source['content_snippet'])
                
                response_timestamp = datetime.now().strftime("%H:%M:%S")
                st.caption(f" {response_timestamp} |  {result['response_time']:.2f}s")
                
                # Add assistant message to history
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": result["response"],
                    "sources": result["sources"],
                    "timestamp": response_timestamp,
                    "response_time": result['response_time']
                })
        
        # Sample questions
        st.subheader(" Try These Sample Questions")
        col1, col2 = st.columns(2)
        
        sample_questions = [
            "How do I return a product?",
            "What are your shipping options?",
            "How can I reset my password?",
            "What payment methods do you accept?",
            "My device won't turn on, what should I do?",
            "How do I track my order?"
        ]
        
        for i, question in enumerate(sample_questions):
            col = col1 if i % 2 == 0 else col2
            if col.button(f" {question}", key=f"sample_{i}"):
                # Simulate user asking the question
                st.session_state.messages.append({
                    "role": "user",
                    "content": question,
                    "timestamp": datetime.now().strftime("%H:%M:%S")
                })
                st.rerun()
    
    with tab2:
        st.subheader(" Support Analytics Dashboard")
        
        analytics = st.session_state.rag_system.analytics_data
        
        if analytics['queries']:
            # Metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total Queries", len(analytics['queries']))
            
            with col2:
                avg_response_time = np.mean(analytics['response_times']) if analytics['response_times'] else 0
                st.metric("Avg Response Time", f"{avg_response_time:.2f}s")
            
            with col3:
                st.metric("Active Sessions", "1")
            
            with col4:
                resolution_rate = 95  # Simulated
                st.metric("Resolution Rate", f"{resolution_rate}%")
            
            # Topic distribution
            col1, col2 = st.columns(2)
            
            with col1:
                if analytics['topics']:
                    topic_counts = pd.Series(analytics['topics']).value_counts()
                    fig = px.pie(
                        values=topic_counts.values,
                        names=topic_counts.index,
                        title="Query Topics Distribution"
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if analytics['response_times']:
                    fig = px.histogram(
                        x=analytics['response_times'],
                        title="Response Time Distribution",
                        labels={'x': 'Response Time (seconds)', 'y': 'Frequency'}
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            # Recent queries
            st.subheader(" Recent Queries")
            recent_df = pd.DataFrame({
                'Query': analytics['queries'][-10:],
                'Topic': analytics['topics'][-10:],
                'Response Time (s)': [f"{t:.2f}" for t in analytics['response_times'][-10:]]
            })
            st.dataframe(recent_df, use_container_width=True)
            
        else:
            st.info(" No analytics data available yet. Start chatting to see metrics!")
    
    with tab3:
        st.subheader(" Knowledge Base Overview")
        
        # Knowledge base statistics
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Default Documents", "5")
        
        with col2:
            st.metric("Uploaded Documents", len(st.session_state.uploaded_docs))
        
        with col3:
            total_docs = 5 + len(st.session_state.uploaded_docs)
            st.metric("Total Documents", total_docs)
        
        # Document categories
        categories = {
            'Returns': 'Information about return policies and procedures',
            'Shipping': 'Shipping options, costs, and delivery information',
            'Account': 'Account management and login assistance',
            'Technical': 'Technical support and troubleshooting guides',
            'Payment': 'Payment methods and billing information',
            'Uploaded': 'User-uploaded custom documents'
        }
        
        st.subheader("Document Categories")
        for category, description in categories.items():
            with st.expander(f" {category}"):
                st.write(description)
                if category == 'Uploaded' and st.session_state.uploaded_docs:
                    for doc in st.session_state.uploaded_docs:
                        st.write(f"• {doc['name']} ({doc['upload_time']})")
                elif category != 'Uploaded':
                    st.write("• Default knowledge base document included")
        
        # Search functionality
        st.subheader("Search Knowledge Base")
        search_query = st.text_input("Search documents...")
        
        if search_query and st.session_state.rag_system.vector_store:
            with st.spinner("Searching..."):
                docs = st.session_state.rag_system.vector_store.similarity_search(search_query, k=5)
                
                if docs:
                    st.write(f"Found {len(docs)} relevant documents:")
                    for i, doc in enumerate(docs):
                        with st.expander(f" {doc.metadata.get('title', f'Document {i+1}')}"):
                            st.write(f"**Category:** {doc.metadata.get('category', 'Unknown')}")
                            st.write(f"**Content Preview:** {doc.page_content[:300]}...")
                else:
                    st.info("No relevant documents found.")
    
    # Footer
    st.divider()
    st.markdown("""
    **🔧 Technical Implementation:**
    - **RAG Pipeline:** Document retrieval → Context injection → Response generation
    - **Vector Store:** FAISS for efficient similarity search
    - **Embeddings:** OpenAI embeddings (with HuggingFace fallback)
    - **LLM:** GPT-3.5-turbo for response generation
    - **Document Processing:** Support for TXT, PDF, and DOCX files
    """)


if __name__ == "__main__":
    main()